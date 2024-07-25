import sys
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QApplication, QMainWindow,QComboBox, QListWidget, QListWidgetItem, QAbstractItemView, QVBoxLayout, QPushButton, QLabel, QLineEdit, QWidget, QMessageBox, QHBoxLayout
from PyQt5.QtGui import QIcon, QFont
from PyQt5.QtCore import QDate, Qt
from indicaciones import indicaciones
from pictogramas import pictogramas
from docx import Document
from docx.shared import Inches, Pt

def window():
    app = QApplication(sys.argv)
    win = QMainWindow()
    win.setGeometry(200, 200, 1200, 700)
    win.setWindowTitle("Label maker")
    win.setWindowIcon(QIcon("icon.jpg"))

    central_widget = QWidget(win)
    win.setCentralWidget(central_widget)

    # Crear el layout principal
    main_layout = QHBoxLayout(central_widget)

    # Columna izquierda
    left_layout = QVBoxLayout()
    
    lbl_titulo = QLabel("Titulo:")
    left_layout.addWidget(lbl_titulo)
    
    txt_titulo = QLineEdit()
    txt_titulo.setFixedHeight(45)
    txt_titulo.setStyleSheet("font-size: 18px; font-weight: bold;")
    left_layout.addWidget(txt_titulo)

    lbl_lote = QLabel("Lote:")
    left_layout.addWidget(lbl_lote)
    
    txt_lote = QLineEdit()
    left_layout.addWidget(txt_lote)

    lbl_fecha = QLabel("Fecha:")
    left_layout.addWidget(lbl_fecha)
    
    txt_fecha = QLineEdit()
    txt_fecha.setText(QDate.currentDate().toString("dd/MM/yyyy"))
    left_layout.addWidget(txt_fecha)

    lbl_kgs = QLabel("Kgs:")
    left_layout.addWidget(lbl_kgs)
    
    txt_kgs = QLineEdit()
    left_layout.addWidget(txt_kgs)

    #Añadir txt_peligrosidad (selector multiple)

    palabras_advertencia = [
    "",
    "Peligro",
    "Peligro grave",
    "Atención",
    "Advertencia",
    "Precaución"
]

    lbl_advertencia_text = "Advertencia:"
    lbl_advertencia = QLabel(lbl_advertencia_text)
    left_layout.addWidget(lbl_advertencia)

    txt_advertencia = QComboBox()
    
    for advertencia_item in palabras_advertencia:
        txt_advertencia.addItem(advertencia_item)
    
    left_layout.addWidget(txt_advertencia)

    # Añadir txt_contiene (selector múltiple)
    lbl_contiene_text = "Contiene:"
    lbl_contiene = QLabel(lbl_contiene_text)
    left_layout.addWidget(lbl_contiene)

    txt_contiene = QListWidget()
    txt_contiene.setSelectionMode(QAbstractItemView.MultiSelection)
    txt_contiene.setMaximumHeight(100)
    
    contiene_items = ["alcohol bencílico", "3-aminometil-3,5,5-trimetilciclohexilamina"]
    for contiene_item in contiene_items:
        contiene_list_item = QListWidgetItem(contiene_item)
        txt_contiene.addItem(contiene_list_item)
    
    left_layout.addWidget(txt_contiene)


    # Añadir txt_UNnumber (campo de texto)
    lbl_UNnumber_text = "UN Number:"
    lbl_UNnumber = QLabel(lbl_UNnumber_text)
    left_layout.addWidget(lbl_UNnumber)

    txt_UNnumber = QLineEdit()
    left_layout.addWidget(txt_UNnumber)

    # Columna derecha
    right_layout = QVBoxLayout()
    
    lbl_precauciones_text = "Precauciones:"
    lbl_precauciones = QLabel(lbl_precauciones_text)
    right_layout.addWidget(lbl_precauciones)

    txt_precauciones = QListWidget()
    txt_precauciones.setSelectionMode(QAbstractItemView.MultiSelection)
    
    for item in indicaciones:
        list_item = QListWidgetItem(item)
        txt_precauciones.addItem(list_item)
    
    right_layout.addWidget(txt_precauciones)

    # Añadir segundo selector múltiple para imágenes
    lbl_img_text = "Imágenes:"
    lbl_img = QLabel(lbl_img_text)
    right_layout.addWidget(lbl_img)

    txt_img = QListWidget()
    txt_img.setSelectionMode(QAbstractItemView.MultiSelection)
    
    for img_item in pictogramas:
        img_list_item = QListWidgetItem(img_item)
        txt_img.addItem(img_list_item)
    
    right_layout.addWidget(txt_img)




    
    # LOGICA--------------------------------------------------------------------------------------------------------------------

    def clicked(parent):
        print("button clicked")
        
        #doc = Document()
        
        #p = doc.add_paragraph()
        #r = p.add_run()
        #r.add_text('Good Morning every body,This is my ')
        #r.add_picture('icon.jpg',width=Inches(4.0), height=Inches(.7))
        #r.add_text(' do you like it?')
        #p1 = doc.add_paragraph("First Paragraph.")
        #p2 = doc.add_paragraph("Second Paragraph.")

        #table = doc.add_table(rows=3, cols=2)
        #table.cell(0,0).text = "0,0"
        #table.cell(0,1).text = "0,1"
        #table.cell(1,0).text = "1,0"
        #table.cell(1,1).text = "1,1"
        #doc.save('tabletest.docx')
        print(indicaciones["H200"])
        msg = QMessageBox(parent)
        msg.setIcon(QMessageBox.Information)
        msg.setText("Datos guardados correctamente.")
        msg.setWindowTitle("Guardar")
        msg.setStandardButtons(QMessageBox.Ok)
        msg.exec_()

        

        
    def loadClicked(parent):
        print("load clicked")
        #print(txt_titulo.text() + " " + txt_lote.text())
        document = Document("og3.docx")
        #document.paragraphs[0].text = "YEAH"
        #print(document.paragraphs[0].text)
        #fullText = []
        #for para in document.paragraphs:
        #    print("hi")
        #    fullText.append(para.text)
        #print('\n'.join(fullText))


        table = document.tables[0]

        # Title Cell (1,0)

        titlerun = table.cell(1,0).add_paragraph().add_run()
        font = titlerun.font
        font.name = "Arial Black"
        font.size = Pt(18)
        titlerun.add_text(txt_titulo.text())

        # Title Cell (1,0) second part
        titlerun2 = table.cell(1,0).add_paragraph().add_run()
        font2 = titlerun2.font
        font2.name = "Arial Black"
        font2.size = Pt(9)
        titlerun2.add_text("Nº LOTE: "+ txt_lote.text() + "         FECHA: " + txt_fecha.text() +"              KG: " + txt_kgs.text() )
        print("done")

        #Indicaciones Cell (2,0)
        this_run = ""
        #Peligrosidad
        if(txt_advertencia.currentText()!=""):
            this_run = table.cell(2,0).add_paragraph().add_run()
            this_run.add_text(txt_advertencia.currentText())
            this_run.font.size = Pt(11)


        #Indicaciones
        
        for selection in txt_precauciones.selectedItems():
            this_run = table.cell(2,0).add_paragraph().add_run()
            this_run.add_text(selection.text() + ": " + indicaciones[selection.text()])
            this_run.font.size = Pt(8)

        #Meter imagenes

        for img_name in txt_img.selectedItems():
            print(img_name.text())
            document.add_picture("pictogramas/"+ img_name.text()+ ".png", width=Inches(1.26))

        #UN number Cell (3,1)
        this_run = table.cell(3,1).add_paragraph().add_run()
        this_run.add_text("  UN" + txt_UNnumber.text())
        this_run.font.size = Pt(18)
        this_run.font.name = "Arial Black"

        #Contiene Cell (3,0)
        this_run = table.cell(3,0).add_paragraph().add_run()
        this_run.add_text("Contiene: ")
        this_run.font.size = Pt(8)
        for selection in txt_contiene.selectedItems():
            this_run = table.cell(3,0).add_paragraph().add_run()
            this_run.add_text(selection.text())
            this_run.font.size = Pt(6)

    
        #PopUp ha funcionado

        msg = QMessageBox(parent)
        msg.setIcon(QMessageBox.Information)
        msg.setText("Etiqueta creada")
        msg.setWindowTitle("Aceptar")
        msg.setStandardButtons(QMessageBox.Ok)
        msg.exec_()

        #Guardar documento
        document.save("etiquetas/" + txt_titulo.text() +".docx")


        # Añadir botones Save y Load a la columna izquierda
    #btn_save = QPushButton("Save")
    #btn_save.clicked.connect(lambda: clicked(win))
    #left_layout.addWidget(btn_save)

    btn_load = QPushButton("Crear Etiqueta")
    btn_load.clicked.connect(lambda: loadClicked(win))
    btn_load.setFixedHeight(100)
    fuente_btn = QFont()
    fuente_btn.setPointSize(20)
    btn_load.setFont(fuente_btn)

    btn_load.setStyleSheet("""
            QPushButton {
                border: 2px solid black;  /* Borde de 2px de grosor y color negro */
                border-radius: 10px;  /* Bordes redondeados */
                padding: 10px;  /* Espaciado interno */
            }
            QPushButton:hover {
                border: 2px solid red;  /* Cambiar el color del borde cuando se pasa el mouse */
            }
            QPushButton:pressed {
                border: 2px solid blue;  /* Cambiar el color del borde cuando se presiona el botón */
            }
        """)
    
    left_layout.addWidget(btn_load)

        # Agregar los layouts izquierdo y derecho al layout principal
    main_layout.addLayout(left_layout)
    main_layout.addLayout(right_layout)

    win.show()
    sys.exit(app.exec_())
    


window()